"""
Resume Export Utilities - Handles formatting and export of resumes in different formats.

IMPORTANT: PDF generation requires wkhtmltopdf to be installed on your system.
Download it from: https://wkhtmltopdf.org/downloads.html

For Windows: Run the installer and make sure it's added to your PATH
For macOS: brew install wkhtmltopdf
For Ubuntu/Debian: sudo apt-get install wkhtmltopdf
"""
import io
import os
import sys
import subprocess
from pathlib import Path
import jinja2
import pdfkit
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# Disable WeasyPrint - We'll use only wkhtmltopdf
WEASYPRINT_AVAILABLE = False

# Define constants for styling
STYLE_ATS = "ATS-Friendly"
STYLE_MODERN = "Modern"
STYLE_PROFESSIONAL = "Professional"

# Create a directory for templates if it doesn't exist
template_dir = Path("templates")
if not template_dir.exists():
    template_dir.mkdir()

# ---------------------------------
# Global Configuration
# ---------------------------------

# Custom path to wkhtmltopdf - change this if wkhtmltopdf is installed but not in PATH
# For example: WKHTMLTOPDF_PATH = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"
# Check installation page: https://wkhtmltopdf.org/downloads.html
WKHTMLTOPDF_PATH = "C:\\Program Files\\wkhtmltopdf\\bin\\wkhtmltopdf.exe"  # Path to your wkhtmltopdf installation

# Flag to disable problematic imports for simpler deployment
ENABLE_PDF_GENERATION = True  # PDF generation is now enabled

# Force update of templates on startup
def force_update_templates():
    """Force update the template files with the current template definitions."""
    print("Forcing update of template files...")
    for style, filename in template_files.items():
        file_path = template_dir / filename
        template_content = None
        if style == STYLE_ATS:
            template_content = ATS_TEMPLATE
        elif style == STYLE_MODERN:
            template_content = MODERN_TEMPLATE
        elif style == STYLE_PROFESSIONAL:
            template_content = PROFESSIONAL_TEMPLATE
        
        if template_content:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(template_content)
            print(f"Updated template: {filename}")

# Check if wkhtmltopdf is installed and accessible
def is_wkhtmltopdf_installed():
    """Check if wkhtmltopdf is installed and accessible in the PATH."""
    try:
        # Use custom path if provided
        wkhtmltopdf_cmd = WKHTMLTOPDF_PATH if WKHTMLTOPDF_PATH else 'wkhtmltopdf'
        
        # Try to execute wkhtmltopdf to check if it's in the PATH
        if sys.platform.startswith('win'):
            # On Windows, we need to suppress the console window
            startupinfo = subprocess.STARTUPINFO()
            startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
            process = subprocess.Popen(
                [wkhtmltopdf_cmd, '-V'], 
                stdout=subprocess.PIPE, 
                stderr=subprocess.PIPE,
                startupinfo=startupinfo
            )
        else:
            # On Unix-like systems
            process = subprocess.Popen(
                [wkhtmltopdf_cmd, '-V'], 
                stdout=subprocess.PIPE, 
                stderr=subprocess.PIPE
            )
        
        # Wait for the process to complete and get output
        stdout, stderr = process.communicate()
        
        # If the command ran successfully, wkhtmltopdf is installed
        return process.returncode == 0
    except (subprocess.SubprocessError, FileNotFoundError):
        # If there was an error running the command, wkhtmltopdf is not installed
        return False

# ---------------------------------
# HTML Templates for PDF Generation
# ---------------------------------

# ATS-Friendly Template - Simple, clean, no fancy formatting that could confuse ATS systems
ATS_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ name }}</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            font-size: 12px;
            line-height: 1.5;
            margin: 0.75in;
        }
        .header {
            margin-bottom: 15px;
            border-bottom: 1px solid #000;
            padding-bottom: 8px;
        }
        .header h1 {
            font-size: 18px;
            margin: 0 0 5px 0;
            border-bottom: none;
        }
        .contact-info {
            font-size: 11px;
            margin: 5px 0;
        }
        .contact-item {
            display: inline-block;
            margin-right: 12px;
        }
        h1, h2, h3 {
            margin-top: 12px;
            margin-bottom: 6px;
        }
        h1 {
            font-size: 18px;
            border-bottom: 1px solid #000;
            padding-bottom: 5px;
        }
        h2 {
            font-size: 16px;
            text-transform: uppercase;
        }
        .company-name {
            margin-top: -5px;
            margin-bottom: 8px;
            font-style: normal;
        }
        .project-title {
            font-weight: bold;
            font-size: 12px;  /* match bullet size */
            margin-top: 6px;
            margin-bottom: 2px;
        }

        .section {
            margin-bottom: 15px;
        }
        .bullet-list {
            margin-left: 15px;
            list-style-type: disc;
        }
        .experience-entry {
            margin-bottom: 10px;
        }
    </style>
</head>
<body>
  <div class="header">
    <h1>{{ name }}</h1>
    {% if title or location or phone or email %}
    <p>
      {% if title %}{{ title }}{% endif %}
      {% if title and location %} | {% endif %}
      {% if location %}{{ location }}{% endif %}
      {% if phone %} | {{ phone }}{% endif %}
      {% if email %} | {{ email }}{% endif %}
    </p>
    {% endif %}
    {% if linkedin or github or website %}
    <p>
      {% if linkedin %}<a href="{{ linkedin }}">{{ linkedin }}</a>{% endif %}
      {% if github %} | <a href="{{ github }}">{{ github }}</a>{% endif %}
      {% if website %} | <a href="{{ website }}">{{ website }}</a>{% endif %}
    </p>
    {% endif %}
  </div>

  {{ content | safe }}
</body>
</html>
"""

# Modern Template - More whitespace, clean design, some color accents
MODERN_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ name }}</title>
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            font-size: 12px;
            line-height: 1.6;
            margin: 0.75in;
            color: #333;
        }
        .header {
            margin-bottom: 20px;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
        }
        .header h1 {
            font-size: 22px;
            margin: 0 0 5px 0;
            color: #2c3e50;
            border-bottom: none;
        }
        .contact-info {
            font-size: 11px;
            margin: 5px 0;
            color: #555;
        }
        .contact-item {
            display: inline-block;
            margin-right: 10px;
        }
        h1, h2, h3 {
            margin-top: 15px;
            margin-bottom: 8px;
            color: #2c3e50;
        }
        h1 {
            font-size: 20px;
            border-bottom: 2px solid #3498db;
            padding-bottom: 8px;
        }
        h2 {
            font-size: 16px;
            text-transform: uppercase;
            letter-spacing: 1px;
        }
        .company-name {
            margin-top: -4px;
            margin-bottom: 8px;
            color: #555;
            font-style: italic;
        }
        .project-title {
            font-weight: bold;
            font-size: 12px;
            margin-top: 6px;
            margin-bottom: 2px;
            color: #2c3e50;  /* or match your bullet font color */
        }

        .section {
            margin-bottom: 20px;
        }
        .bullet-list {
            margin-left: 18px;
            list-style-type: square;
        }
        .experience-entry {
            margin-bottom: 15px;
            padding-bottom: 10px;
            border-bottom: 1px dotted #ddd;
        }
        h3 {
            font-size: 14px;
            font-weight: bold;
            margin-top: 8px;
            margin-bottom: 4px;
            color: #2c3e50;
        }

    </style>
</head>
<body>
  <div class="header">
    <h1>{{ name }}</h1>
    {% if title or location or phone or email %}
    <p>
      {% if title %}{{ title }}{% endif %}
      {% if title and location %} | {% endif %}
      {% if location %}{{ location }}{% endif %}
      {% if phone %} | {{ phone }}{% endif %}
      {% if email %} | {{ email }}{% endif %}
    </p>
    {% endif %}
    {% if linkedin or github or website %}
    <p>
      {% if linkedin %}<a href="{{ linkedin }}">{{ linkedin }}</a>{% endif %}
      {% if github %} | <a href="{{ github }}">{{ github }}</a>{% endif %}
      {% if website %} | <a href="{{ website }}">{{ website }}</a>{% endif %}
    </p>
    {% endif %}
  </div>

  {{ content | safe }}
</body>
</html>
"""

# Professional Template - Traditional, serif fonts, structured layout
PROFESSIONAL_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ name }}</title>
    <style>
        body {
            font-family: Georgia, 'Times New Roman', Times, serif;
            font-size: 12px;
            line-height: 1.5;
            margin: 0.75in;
            color: #222;
        }
        .header {
            text-align: center;
            margin-bottom: 15px;
            border-bottom: 1px solid #444;
            padding-bottom: 10px;
        }
        .header h1 {
            font-size: 22px;
            margin: 0 0 5px 0;
            border-bottom: none;
            text-transform: uppercase;
            letter-spacing: 1px;
        }
        .contact-info {
            font-size: 11px;
            margin: 5px 0;
        }
        .contact-item {
            display: inline-block;
            margin: 0 8px;
        }
        h1, h2, h3 {
            margin-top: 14px;
            margin-bottom: 7px;
            font-family: Cambria, Cochin, Georgia, Times, 'Times New Roman', serif;
        }
        h1 {
            font-size: 18px;
            border-bottom: 1px solid #444;
            padding-bottom: 6px;
            text-align: center;
        }
        h2 {
            font-size: 16px;
            text-transform: uppercase;
            text-align: center;
        }
        .company-name {
            margin-top: -3px;
            margin-bottom: 10px;
            font-style: italic;
            text-align: center;
        }

        .project-title {
            font-weight: bold;
            font-size: 12px;
            margin-top: 6px;
            margin-bottom: 2px;
            font-family: Georgia, 'Times New Roman', Times, serif;
        }

        .section {
            margin-bottom: 18px;
        }
        .bullet-list {
            margin-left: 20px;
            list-style-type: disc;
        }
        .experience-entry {
            margin-bottom: 12px;
        }
    </style>
</head>
<body>
  <div class="header">
    <h1>{{ name }}</h1>
    {% if title or location or phone or email %}
    <p>
      {% if title %}{{ title }}{% endif %}
      {% if title and location %} | {% endif %}
      {% if location %}{{ location }}{% endif %}
      {% if phone %} | {{ phone }}{% endif %}
      {% if email %} | {{ email }}{% endif %}
    </p>
    {% endif %}
    {% if linkedin or github or website %}
    <p>
      {% if linkedin %}<a href="{{ linkedin }}">{{ linkedin }}</a>{% endif %}
      {% if github %} | <a href="{{ github }}">{{ github }}</a>{% endif %}
      {% if website %} | <a href="{{ website }}">{{ website }}</a>{% endif %}
    </p>
    {% endif %}
  </div>

  {{ content | safe }}
</body>
</html>
"""

# Save templates to files
template_files = {
    STYLE_ATS: "ats_template.html",
    STYLE_MODERN: "modern_template.html",
    STYLE_PROFESSIONAL: "professional_template.html"
}

# Update template files (ensure they have the latest content)
force_update_templates()

# ---------------------------------
# Formatting Helper Functions
# ---------------------------------

def format_text_to_html(resume_text):
    """
    Convert plain text resume with section headers into structured HTML.
    Detects sections, bullet points, and formats them appropriately.
    
    Args:
        resume_text (str): The raw resume text
        
    Returns:
        str: HTML formatted resume content
    """
    # Split the resume into lines
    lines = resume_text.strip().split('\n')
    
    # Initialize variables
    html_output = []
    current_section = None
    in_bullet_list = False
    in_experience_entry = False
    
    # Initialize job position tracking
    previous_line_was_job_title = False
    company_found = False
    
    # Common job title words for detection
    job_title_keywords = [
        'Analyst', 'Engineer', 'Manager', 'Director', 'Specialist', 'Associate', 
        'Consultant', 'Developer', 'Accountant', 'Coordinator', 'Assistant',
        'Supervisor', 'Representative', 'Administrator', 'Designer', 'Programmer',
        'Budget'
    ]
    
    for i, line in enumerate(lines):
        line = line.strip()
        
        # Skip empty lines
        if not line:
            continue
        
        # Detect section headers (all caps, or with dashes/underscores)
        if line.isupper() or line.startswith('---') or all(c in line for c in ['=', '=']):
            # Close any open sections and lists
            if in_bullet_list:
                html_output.append('</ul>')
                in_bullet_list = False
            
            # Start a new section
            clean_line = line.replace('-', '').replace('=', '').strip()
            html_output.append(f'<h2>{clean_line}</h2>')
            html_output.append('<div class="section">')
            if current_section:
                html_output.append('</div>')  # Close previous section
            current_section = clean_line
            previous_line_was_job_title = False
            company_found = False
            continue
        
        # Special processing for experience section
        if current_section and current_section.upper() in ["EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT", "PROFESSIONAL EXPERIENCE"]:
            
            # Check for position with special bullets (□, ■, ◆, 📌)
            special_bullet_match = re.match(r'^([□■◆📌]\s*)(.*?)$', line)
            if special_bullet_match:
                # This is a job title line with a special bullet
                if in_bullet_list:
                    html_output.append('</ul>')
                    in_bullet_list = False
                
                if in_experience_entry:
                    html_output.append('</div>')  # Close previous experience entry
                
                bullet, position_text = special_bullet_match.groups()
                html_output.append('<div class="experience-entry">')
                html_output.append(f'<h3>{position_text}</h3>')
                in_experience_entry = True
                previous_line_was_job_title = True
                company_found = False
                continue
            
            # Check for date pattern directly indicating a job title line
            date_pattern_match = re.search(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\'\d{2}\s*[-–—]\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\'\d{2}|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\'\d{2}\s*[-–—]\s*Present', line, re.IGNORECASE)
            
            # Check for a job title format (Position at Company)
            at_company_pattern = re.search(r'(.*?)\s+at\s+(.*?)$', line, re.IGNORECASE)
            
            # If a line contains a job title keyword followed by date pattern, it's likely a job title
            job_title_with_date = False
            if any(keyword in line for keyword in job_title_keywords) and re.search(r'\b(19|20)\d{2}\s*[-–—]\s*(19|20)\d{2}|\b(19|20)\d{2}\s*[-–—]', line):
                job_title_with_date = True
            
            if date_pattern_match or at_company_pattern or job_title_with_date:
                # Handle as job title
                if in_bullet_list:
                    html_output.append('</ul>')
                    in_bullet_list = False
                
                if in_experience_entry:
                    html_output.append('</div>')  # Close previous experience entry
                
                html_output.append('<div class="experience-entry">')
                html_output.append(f'<h3>{line}</h3>')
                in_experience_entry = True
                previous_line_was_job_title = True
                company_found = False
                continue
            
            # Handle company name after job title
            if previous_line_was_job_title and not company_found:
                # If the next line doesn't look like a bullet point and is relatively short, it's likely a company name
                if (not line.startswith(('•', '-', '*')) and 
                    not re.match(r'^\d+\.', line) and 
                    len(line) < 100):
                    
                    # Handle company with "at" format (e.g., "at Company Name")
                    if line.lower().startswith('at '):
                        company_name = line[3:].strip()  # Remove 'at ' prefix
                    else:
                        company_name = line
                    
                    html_output.append(f'<p class="company-name"><strong>{company_name}</strong></p>')
                    previous_line_was_job_title = False
                    company_found = True
                    continue
        
        # Handle regular bullet points (lines starting with •, -, *, or numbers)
        if line.startswith('•') or line.startswith('-') or line.startswith('*') or (line.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.'))):
            # If the previous line was a job title and no company was found, this isn't a real bullet point yet
            if previous_line_was_job_title and not company_found:
                # We'll assume no company name was provided, so continue to bullet points
                company_found = True
                previous_line_was_job_title = False
                
            if not in_bullet_list:
                html_output.append('<ul class="bullet-list">')
                in_bullet_list = True
            
            # Clean the bullet character and add the item
            bullet_text = line.lstrip('•-*123456789. \t')
            html_output.append(f'<li>{bullet_text}</li>')
            continue
        
        # Regular paragraph text
        if in_bullet_list:
            html_output.append('</ul>')
            in_bullet_list = False
        
        html_output.append(f'<p>{line}</p>')
        previous_line_was_job_title = False
    
    # Close any open tags
    if in_bullet_list:
        html_output.append('</ul>')
    if in_experience_entry:
        html_output.append('</div>')
    if current_section:
        html_output.append('</div>')
    
    return '\n'.join(html_output)

# ---------------------------------
# Export Helper Functions
# ---------------------------------

def generate_pdf(resume_text, style=STYLE_ATS, name="Resume", email=None, phone=None, location=None, linkedin=None, github=None, website=None):
    """
    Generate a PDF version of the resume using the specified style template.
    
    Args:
        resume_text (str): Plain text resume content
        style (str): Template style to use (ATS, Modern, Professional)
        name (str): Name for the resume
        email (str, optional): Email address
        phone (str, optional): Phone number
        location (str, optional): Location (city, state)
        linkedin (str, optional): LinkedIn profile URL
        github (str, optional): GitHub profile URL
        website (str, optional): Personal website URL
        
    Returns:
        bytes: The PDF file as bytes
    """
    # Check if PDF generation is disabled
    if not ENABLE_PDF_GENERATION:
        raise Exception(
            "PDF generation is disabled. To enable it:\n"
            "1. Install wkhtmltopdf from https://wkhtmltopdf.org/downloads.html\n"
            "2. Set ENABLE_PDF_GENERATION = True in resume_export.py\n"
            "3. If needed, set WKHTMLTOPDF_PATH to the executable path"
        )
    
    # Print debugging info for contact info
    print(f"PDF Generation - Contact Info:")
    print(f"  Name: {name}")
    print(f"  Email: {email}")
    print(f"  Phone: {phone}")
    print(f"  Location: {location}")
    print(f"  LinkedIn: {linkedin}")
    print(f"  GitHub: {github}")
    print(f"  Website: {website}")
    
    # Get the appropriate template based on style
    template_file = template_files.get(style, template_files[STYLE_ATS])
    template_path = template_dir / template_file
    
    # Format text to HTML
    html_content = format_text_to_html(resume_text)
    
    # Load and render the template
    with open(template_path, 'r', encoding='utf-8') as f:
        template_text = f.read()
    
    # Create a Jinja2 template and render it
    template = jinja2.Template(template_text)
    rendered_html = template.render(
        name=name if name else "Resume",
        content=html_content,
        email=email if email else None,
        phone=phone if phone else None,
        location=location if location else None,
        linkedin=linkedin if linkedin else None,
        github=github if github else None,
        website=website if website else None
    )
    
    # Save the rendered HTML to a temporary file
    temp_html = template_dir / "temp_resume.html"
    with open(temp_html, 'w', encoding='utf-8') as f:
        f.write(rendered_html)
    
    # Write debug HTML to a file for inspection
    debug_html = template_dir / "debug_resume.html"
    with open(debug_html, 'w', encoding='utf-8') as f:
        f.write(rendered_html)
    
    # Convert HTML to PDF
    pdf_bytes = io.BytesIO()
    
    # First check if wkhtmltopdf is installed
    wkhtmltopdf_available = is_wkhtmltopdf_installed()
    
    if wkhtmltopdf_available:
        # Try using pdfkit (wkhtmltopdf)
        try:
            # Options for PDF generation
            options = {
                'page-size': 'Letter',
                'margin-top': '0.5in',
                'margin-right': '0.5in',
                'margin-bottom': '0.5in',
                'margin-left': '0.5in',
                'encoding': 'UTF-8',
                'quiet': '',
                'print-media-type': '',
                'no-background': '',
                'enable-local-file-access': ''
            }
            
            # Configure pdfkit with custom path if provided
            config = None
            if WKHTMLTOPDF_PATH:
                config = pdfkit.configuration(wkhtmltopdf=WKHTMLTOPDF_PATH)
            
            # Try direct conversion to bytes
            if config:
                pdf_data = pdfkit.from_file(str(temp_html), False, options=options, configuration=config)
            else:
                pdf_data = pdfkit.from_file(str(temp_html), False, options=options)
            
            pdf_bytes.write(pdf_data)
        except Exception as e:
            print(f"Error using pdfkit: {e}")
            # Try fallback method for pdfkit
            try:
                output_path = template_dir / "temp_resume.pdf"
                
                if config:
                    pdfkit.from_file(str(temp_html), str(output_path), options=options, configuration=config)
                else:
                    pdfkit.from_file(str(temp_html), str(output_path), options=options)
                
                with open(output_path, 'rb') as f:
                    pdf_bytes.write(f.read())
                # Clean up temporary PDF file
                if os.path.exists(str(output_path)):
                    os.remove(str(output_path))
            except Exception as e:
                print(f"PDF generation failed: {e}")
                raise Exception("PDF generation failed. Check that wkhtmltopdf is properly installed.") from e
    else:
        raise Exception(
            "wkhtmltopdf is not installed or not found in PATH.\n"
            "Please install it from https://wkhtmltopdf.org/downloads.html\n"
            "Or specify the correct path in resume_export.py (WKHTMLTOPDF_PATH)"
        )
    
    # Clean up temporary HTML file
    if os.path.exists(str(temp_html)):
        os.remove(str(temp_html))
    
    pdf_bytes.seek(0)
    return pdf_bytes.getvalue()

def generate_docx(resume_text, style=STYLE_ATS, name="Resume", email=None, phone=None, location=None, linkedin=None, github=None, website=None):
    """
    Generate a Word document (.docx) version of the resume using the specified style.
    
    Args:
        resume_text (str): Plain text resume content
        style (str): Template style to use (ATS, Modern, Professional)
        name (str): Name for the resume
        email (str, optional): Email address
        phone (str, optional): Phone number
        location (str, optional): Location (city, state)
        linkedin (str, optional): LinkedIn profile URL
        github (str, optional): GitHub profile URL
        website (str, optional): Personal website URL
        
    Returns:
        bytes: The DOCX file as bytes
    """
    # Print debugging info for contact info
    print(f"DOCX Generation - Contact Info:")
    print(f"  Name: {name}")
    print(f"  Email: {email}")
    print(f"  Phone: {phone}")
    print(f"  Location: {location}")
    print(f"  LinkedIn: {linkedin}")
    print(f"  GitHub: {github}")
    print(f"  Website: {website}")
    
    # Create a new Document
    doc = Document()
    
    # Set up document styles based on the selected style
    if style == STYLE_ATS:
        # ATS-Friendly - Simple, no frills
        font_name = "Arial"
        heading_color = None  # Default black
        bullet_type = 'bullet'  # Standard bullet
    elif style == STYLE_MODERN:
        # Modern - Sleek and contemporary
        font_name = "Calibri"
        heading_color = '2c3e50'  # Dark blue
        bullet_type = 'square'  # Square bullets
    else:  # Professional
        # Professional - Traditional, serif
        font_name = "Times New Roman"
        heading_color = None  # Default black
        bullet_type = 'bullet'  # Standard bullet
    
    # Set up document margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
    
    # Ensure name has a value
    if not name:
        print("⚠️ Warning: No name detected for resume header.")
        name = "Your Name"

    
    # Add the header with name and contact information
    # Name
    heading = doc.add_heading(name, level=0)
    if style == STYLE_PROFESSIONAL:
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in heading.runs:
        run.font.name = font_name
        run.font.size = Pt(16)
    
    # Contact Information
    contact_info = []
    if email:
        contact_info.append(email)
    if phone:
        contact_info.append(phone)
    if location:
        contact_info.append(location)
        
    # Add LinkedIn, GitHub, Website
    links_info = []
    if linkedin:
        links_info.append(linkedin)
    if github:
        links_info.append(github)
    if website:
        links_info.append(website)
    
    # Add contact information line
    if contact_info:
        p = doc.add_paragraph()
        if style == STYLE_PROFESSIONAL:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run(" | ".join(contact_info))
        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(10)
    
    # Add links line
    if links_info:
        p = doc.add_paragraph()
        if style == STYLE_PROFESSIONAL:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run(" | ".join(links_info))
        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(10)
    
    # Add a spacer
    doc.add_paragraph()
    
    # Split the resume into lines
    lines = resume_text.strip().split('\n')
    
    # Process each line
    in_bullet_list = False
    current_section = None
    previous_line_was_job_title = False
    company_found = False
    
    job_title_keywords = [
        'Analyst', 'Engineer', 'Manager', 'Director', 'Specialist', 'Associate', 
        'Consultant', 'Developer', 'Accountant', 'Coordinator', 'Assistant',
        'Supervisor', 'Representative', 'Administrator', 'Designer', 'Programmer',
        'Budget'
    ]
    
    for line in lines:
        line = line.strip()
        
        # Skip empty lines
        if not line:
            continue
        
        # Detect section headers (all caps, or with dashes/underscores)
        if line.isupper() or line.startswith('---') or all(c in line for c in ['=', '=']):
            # Add a section header
            clean_line = line.replace('-', '').replace('=', '').strip()
            heading = doc.add_heading(clean_line, level=1)
            
            # Customize heading based on style
            for run in heading.runs:
                run.font.name = font_name
                if style == STYLE_PROFESSIONAL:
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
            current_section = clean_line
            in_bullet_list = False
            previous_line_was_job_title = False
            company_found = False
            continue
        
        # Special processing for experience section
        if current_section and current_section.upper() in ["EXPERIENCE", "WORK EXPERIENCE", "EMPLOYMENT", "PROFESSIONAL EXPERIENCE"]:
            
            # Check for position with special bullets (□, ■, ◆, 📌)
            special_bullet_match = re.match(r'^([□■◆📌]\s*)(.*?)$', line)
            if special_bullet_match:
                # This is a job title line with a special bullet
                bullet, position_text = special_bullet_match.groups()
                heading = doc.add_heading(position_text, level=2)
                
                # Customize heading based on style
                for run in heading.runs:
                    run.font.name = font_name
                    run.font.size = Pt(12)
                    if style == STYLE_PROFESSIONAL:
                        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                previous_line_was_job_title = True
                company_found = False
                in_bullet_list = False
                continue
            
            # Check for date pattern directly indicating a job title line
            date_pattern_match = re.search(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\'\d{2}\s*[-–—]\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\'\d{2}|\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\'\d{2}\s*[-–—]\s*Present', line, re.IGNORECASE)
            
            # Check for a job title format (Position at Company)
            at_company_pattern = re.search(r'(.*?)\s+at\s+(.*?)$', line, re.IGNORECASE)
            
            # If a line contains a job title keyword followed by date pattern, it's likely a job title
            job_title_with_date = False
            if any(keyword in line for keyword in job_title_keywords) and re.search(r'\b(19|20)\d{2}\s*[-–—]\s*(19|20)\d{2}|\b(19|20)\d{2}\s*[-–—]', line):
                job_title_with_date = True
            
            if date_pattern_match or at_company_pattern or job_title_with_date:
                # Add a job title heading
                heading = doc.add_heading(line, level=2)
                
                # Customize heading based on style
                for run in heading.runs:
                    run.font.name = font_name
                    run.font.size = Pt(12)
                    if style == STYLE_PROFESSIONAL:
                        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                previous_line_was_job_title = True
                company_found = False
                in_bullet_list = False
                continue
            
            # Handle company name after job title
            if previous_line_was_job_title and not company_found:
                if (not line.startswith(('•', '-', '*')) and 
                    not re.match(r'^\d+\.', line) and 
                    len(line) < 100):
                    
                    # Handle company with "at" format (e.g., "at Company Name")
                    if line.lower().startswith('at '):
                        company_name = line[3:].strip()  # Remove 'at ' prefix
                    else:
                        company_name = line
                    
                    # Add the company name as italicized paragraph
                    p = doc.add_paragraph()
                    p.add_run(company_name).italic = True
                    
                    for run in p.runs:
                        run.font.name = font_name
                        run.font.size = Pt(11)
                        if style == STYLE_PROFESSIONAL:
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    previous_line_was_job_title = False
                    company_found = True
                    continue
        
        # Handle bullet points
        if line.startswith('•') or line.startswith('-') or line.startswith('*') or (line.strip().startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.'))):
            # If the previous line was a job title and no company was found, this isn't a real bullet point yet
            if previous_line_was_job_title and not company_found:
                # We'll assume no company name was provided, so continue to bullet points
                company_found = True
                previous_line_was_job_title = False
            
            # Add a bullet point
            bullet_text = line.lstrip('•-*123456789. \t')
            p = doc.add_paragraph(bullet_text)
            p.style = 'List Bullet'
            
            # Format the bullet point
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(11)
            
            in_bullet_list = True
            continue
        
        # Regular paragraph text
        p = doc.add_paragraph(line)
        
        # Format paragraph
        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(11)
    
    # Save the document to a byte stream
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()

# List of available styles
def get_available_styles():
    """Return a list of available resume styles."""
    return [STYLE_ATS, STYLE_MODERN, STYLE_PROFESSIONAL] 